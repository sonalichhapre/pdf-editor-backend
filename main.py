from __future__ import annotations

import io
import os
import subprocess

# Auto-install LibreOffice if missing
if not os.path.exists('/usr/bin/soffice'):
    subprocess.run(['apt-get', 'update', '-y'], check=False)
    subprocess.run(['apt-get', 'install', '-y', 'libreoffice'], check=False)

os.environ.setdefault('LIBREOFFICE_PATH', '/usr/bin/soffice')
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from collections import defaultdict
from datetime import date

from fastapi import FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from starlette.background import BackgroundTask

app = FastAPI(title="PDF & Word Converter")

# Flow: Accept upload → Process (LibreOffice/pypdf) → Store in temp dir → Send file → Delete temp (BackgroundTask)

# ---------- Freemium rate limit (5 per day for free; bypass with PRO_API_KEY) ----------
FREE_DAILY_LIMIT = int(os.getenv("FREE_DAILY_LIMIT", "999999"))  # Effectively unlimited (all free for now)
_rate_store: dict[str, dict[str, int]] = defaultdict(lambda: {"date": "", "count": 0})


def _get_client_key(request: Request) -> str:
    forwarded = request.headers.get("X-Forwarded-For")
    if forwarded:
        return forwarded.split(",")[0].strip()
    return request.client.host if request.client else "unknown"


def _check_rate_limit(request: Request) -> None:
    if os.getenv("PRO_API_KEY") and request.headers.get("X-API-Key") == os.getenv("PRO_API_KEY"):
        return
    key = _get_client_key(request)
    today = date.today().isoformat()
    data = _rate_store[key]
    if data["date"] != today:
        data["date"] = today
        data["count"] = 0
    if data["count"] >= FREE_DAILY_LIMIT:
        raise HTTPException(
            status_code=429,
            detail=f"Daily limit reached ({FREE_DAILY_LIMIT} free conversions). Upgrade to unlimited for $4.99/month.",
        )


def _increment_usage(request: Request) -> None:
    if os.getenv("PRO_API_KEY") and request.headers.get("X-API-Key") == os.getenv("PRO_API_KEY"):
        return
    key = _get_client_key(request)
    today = date.today().isoformat()
    data = _rate_store[key]
    if data["date"] != today:
        data["date"] = today
        data["count"] = 0
    data["count"] += 1


# ---------- CORS ----------
_cors_origins = os.getenv(
    "CORS_ALLOW_ORIGINS",
    "http://localhost:5173,http://127.0.0.1:5173,https://getdocease.vercel.app,https://pdf-editor-frontend-cyan.vercel.app",
)
allow_origins = [o.strip() for o in _cors_origins.split(",") if o.strip()]

app.add_middleware(
    CORSMiddleware,
    allow_origins=allow_origins,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------- Conversion engine ----------
# Use unoserver for 2–4x faster conversions (LibreOffice stays running).
# Run: unoserver --port 2003
# Set: USE_UNOSERVER=true
_USE_UNOSERVER = os.getenv("USE_UNOSERVER", "false").lower() in ("true", "1", "yes")
_UNOSERVER_HOST = os.getenv("UNOSERVER_HOST", "127.0.0.1")
_UNOSERVER_PORT = os.getenv("UNOSERVER_PORT", "2003")

# LibreOffice --convert-to format strings
_LO_FORMAT = {
    ".html": "html",
    ".pdf": "pdf:writer_pdf_Export",
    ".docx": "docx:MS Word 2007 XML",
}


def _safe_basename(filename: str) -> str:
    name = os.path.basename(filename or "").strip()
    return name or "upload"


def _run_convert(
    *,
    input_path: Path,
    output_path: Path,
    timeout_s: int = 90,
    infilter: Optional[str] = None,
) -> None:
    """Convert file using unoserver (fast) or direct LibreOffice (slow cold start)."""
    if _USE_UNOSERVER:
        cmd = [
            "unoconvert",
            "--host", _UNOSERVER_HOST,
            "--port", _UNOSERVER_PORT,
            str(input_path),
            str(output_path),
        ]
    else:
        soffice = os.getenv("LIBREOFFICE_PATH", "libreoffice")
        ext = output_path.suffix.lower()
        convert_to = _LO_FORMAT.get(ext, ext.lstrip("."))
        # Isolated profile avoids dconf/DeploymentException crashes in server environments
        profile_dir = Path(tempfile.gettempdir()) / "lo_pdf_editor"
        profile_dir.mkdir(exist_ok=True)
        profile_uri = "file://" + str(profile_dir.resolve()).replace("\\", "/")
        cmd = [
            soffice,
            "--headless",
            "--nologo",
            "--nolockcheck",
            "--nodefault",
            "--norestore",
            "--invisible",
            f"-env:UserInstallation={profile_uri}",
        ]
        if infilter:
            cmd.extend([f"-infilter={infilter}"])
        cmd.extend([
            "--convert-to",
            convert_to,
            "--outdir",
            str(output_path.parent),
            str(input_path),
        ])

    env = os.environ.copy()
    env["SAL_USE_VCLPLUGIN"] = "gen"
    env["HOME"] = "/tmp"
    env["DISPLAY"] = ":99"
    env["DBUS_SESSION_BUS_ADDRESS"] = "/dev/null"

    try:
        process = subprocess.run(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=timeout_s,
            check=False,
            env=env,
        )
    except FileNotFoundError as e:
        binary = "unoconvert" if _USE_UNOSERVER else soffice
        raise HTTPException(
            status_code=500,
            detail=f"Conversion binary not found ({binary}). "
            + ("Install unoserver and run 'unoserver --port 2003'." if _USE_UNOSERVER else "Set LIBREOFFICE_PATH."),
        ) from e
    except subprocess.TimeoutExpired as e:
        raise HTTPException(
            status_code=500,
            detail=f"Conversion timed out after {timeout_s}s.",
        ) from e

    stderr = (process.stderr or b"").decode(errors="replace").strip()
    stdout = (process.stdout or b"").decode(errors="replace").strip()
    print(f"CMD: {cmd}", flush=True)
    print(f"RETURNCODE: {process.returncode}", flush=True)
    print(f"STDOUT: {stdout}", flush=True)
    print(f"STDERR: {stderr}", flush=True)
    if process.returncode != 0:
        detail = stderr or stdout or "Conversion failed"
        raise HTTPException(status_code=500, detail=detail)


def _reduce_pdf_to_size(pdf_path: Path, target_bytes: int) -> Path:
    """Compress PDF using Ghostscript until under target size. Returns path to compressed file."""
    gs = os.getenv("GHOSTSCRIPT_PATH", "gs")
    current = pdf_path
    stem = pdf_path.stem

    def run_gs(input_p: Path, output_p: Path, extra_args: list) -> bool:
        try:
            subprocess.run(
                [
                    gs,
                    "-sDEVICE=pdfwrite",
                    "-dCompatibilityLevel=1.4",
                    "-dNOPAUSE",
                    "-dQUIET",
                    "-dBATCH",
                    *extra_args,
                    f"-sOutputFile={output_p}",
                    str(input_p),
                ],
                capture_output=True,
                timeout=90,
                check=True,
            )
            return output_p.exists()
        except (FileNotFoundError, subprocess.CalledProcessError):
            return False

    # 1. Try PDFSETTINGS presets (screen = smallest)
    for setting in ("/screen", "/ebook", "/printer"):
        out_path = pdf_path.parent / f"{stem}_c{setting.replace('/', '')}.pdf"
        if run_gs(current, out_path, [f"-dPDFSETTINGS={setting}"]) and out_path.stat().st_size < current.stat().st_size:
            current = out_path
        if current.stat().st_size <= target_bytes:
            return current

    # 2. Iteratively try lower custom resolutions until target met
    for dpi in (72, 50, 36, 24):
        out_path = pdf_path.parent / f"{stem}_dpi{dpi}.pdf"
        extra = [
            "-dColorImageResolution=%d" % dpi,
            "-dGrayImageResolution=%d" % dpi,
            "-dMonoImageResolution=%d" % dpi,
            "-dDownsampleColorImages=true",
            "-dDownsampleGrayImages=true",
            "-dDownsampleMonoImages=true",
        ]
        if run_gs(current, out_path, extra) and out_path.stat().st_size < current.stat().st_size:
            current = out_path
        if current.stat().st_size <= target_bytes:
            return current

    return current


def _reduce_docx_to_size(docx_path: Path, target_bytes: int) -> Path:
    """Compress DOCX by recompressing and resizing images. Returns path to compressed file."""
    try:
        from PIL import Image
    except ImportError:
        return docx_path

    img_ext = (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".tif")
    # Try progressively more aggressive: quality then max dimension
    for max_dim, qualities in [(1200, (85, 70, 50, 30)), (800, (25, 20, 15)), (500, (15, 10))]:
        out_path = docx_path.parent / f"{docx_path.stem}_compressed.docx"
        with zipfile.ZipFile(docx_path, "r") as z_in:
            with zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as z_out:
                for item in z_in.namelist():
                    data = z_in.read(item)
                    if item.lower().endswith(img_ext) and "media/" in item.lower():
                        try:
                            img = Image.open(io.BytesIO(data))
                            if img.mode in ("RGBA", "P"):
                                img = img.convert("RGB")
                            w, h = img.size
                            if max(w, h) > max_dim:
                                ratio = max_dim / max(w, h)
                                new_size = (int(w * ratio), int(h * ratio))
                                img = img.resize(new_size, Image.Resampling.LANCZOS)
                            buf = io.BytesIO()
                            best_data = data
                            for q in qualities:
                                buf.seek(0)
                                buf.truncate()
                                img.save(buf, "JPEG", quality=q, optimize=True)
                                if buf.tell() < len(best_data):
                                    best_data = buf.getvalue()
                            data = best_data
                        except Exception:
                            pass
                    z_out.writestr(item, data)

        if out_path.exists() and out_path.stat().st_size < docx_path.stat().st_size:
            docx_path = out_path
        if docx_path.stat().st_size <= target_bytes:
            return docx_path

    return docx_path


def _add_page_numbers_pdf(
    pdf_path: Path,
    start: int = 1,
    total: Optional[int] = None,
) -> Tuple[Path, int]:
    """Add page numbers to PDF. start=first page number, total=total (default: actual page count). Returns (path, num_pages)."""
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import inch
    from pypdf import PdfReader, PdfWriter

    reader = PdfReader(str(pdf_path))
    if reader.is_encrypted:
        raise ValueError("PDF is password-protected. Please remove the password first.")
    writer = PdfWriter()
    num_pages = len(reader.pages)
    if num_pages == 0:
        raise ValueError("PDF has no pages")

    total_pages = total if total is not None and total > 0 else num_pages

    for i, page in enumerate(reader.pages):
        mediabox = page.mediabox
        width = max(float(mediabox.width or 612), 72)
        height = max(float(mediabox.height or 792), 72)

        page_num = start + i
        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=(width, height))
        c.setFont("Helvetica", 10)
        text = f"Page {page_num} of {total_pages}"
        c.drawRightString(width - 0.5 * inch, 0.5 * inch, text)
        c.save()
        buf.seek(0)
        overlay = PdfReader(buf).pages[0]
        page.merge_page(overlay)
        writer.add_page(page)

    out_path = pdf_path.parent / f"{pdf_path.stem}_numbered.pdf"
    with open(out_path, "wb") as f:
        writer.write(f)
    return out_path, num_pages


def _add_page_numbers_docx(docx_path: Path, start: int = 1, total: Optional[int] = None) -> Path:
    """Add page numbers to Word document footer. Word uses PAGE field (1,2,3...); start/total apply to PDF only."""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    docx_path = Path(docx_path).resolve()
    if not docx_path.exists():
        raise FileNotFoundError(f"Word file not found at {docx_path}")
    doc = Document(str(docx_path))
    for section in doc.sections:
        try:
            footer = section.footer
        except AttributeError:
            continue
        if not footer.paragraphs:
            p = footer.add_paragraph()
        else:
            p = footer.paragraphs[0]
        p.alignment = 1
        run = p.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.extend([fld_begin, instr, fld_sep, fld_end])

    out_path = docx_path.parent / f"{docx_path.stem}_numbered.docx"
    doc.save(str(out_path))
    return out_path


def _remove_watermark_pdf(pdf_path: Path) -> Path:
    """Remove watermarks from PDF using PyMuPDF (annotation-based watermarks)."""
    try:
        import fitz
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="Watermark removal requires PyMuPDF. Install with: pip install pymupdf",
        ) from None

    doc = fitz.open(str(pdf_path))
    for page in doc:
        annots = page.annots(types=[fitz.PDF_ANNOT_WATERMARK]) or []
        for annot in list(annots):
            page.delete_annot(annot)
    out_path = pdf_path.parent / f"{pdf_path.stem}_no_watermark.pdf"
    doc.save(str(out_path))
    doc.close()
    return out_path


def _remove_watermark_docx(docx_path: Path) -> Path:
    """Remove watermarks from Word by clearing header content (common watermark location)."""
    from docx import Document

    doc = Document(str(docx_path))
    for section in doc.sections:
        for header in (section.header, section.first_page_header):
            for p in header.paragraphs:
                p.clear()
        for footer in (section.footer, section.first_page_footer):
            pass
    out_path = docx_path.parent / f"{docx_path.stem}_no_watermark.docx"
    doc.save(str(out_path))
    return out_path


def _find_output(outdir: Path, expected_path: Path, suffix: str) -> Path:
    """Return expected_path if it exists, else find first file with given suffix in outdir (including subdirs)."""
    if expected_path.exists():
        return expected_path
    for p in outdir.rglob("*"):
        if p.is_file() and p.suffix.lower() == suffix.lower():
            return p
    files = [str(p.relative_to(outdir)) for p in outdir.rglob("*") if p.is_file()]
    raise HTTPException(
        status_code=500,
        detail=f"Conversion output ({suffix}) was not generated. Files in output: {files or 'none'}",
    )


def _response_with_tmpdir_cleanup(
    path: Path, *, media_type: str, download_name: str, tmpdir: Path
) -> FileResponse:
    return FileResponse(
        str(path),
        media_type=media_type,
        filename=download_name,
        background=BackgroundTask(shutil.rmtree, str(tmpdir), ignore_errors=True),
    )


@app.get("/health")
@app.head("/health")
def health() -> Dict[str, str]:
    return {"status": "ok"}


# ---------- ADD PAGE NUMBERS ----------
@app.post("/add-page-numbers")
async def add_page_numbers(
    request: Request,
    file: UploadFile = File(...),
    start: Optional[str] = Form(None),
    total: Optional[str] = Form(None),
):
    """Add page numbers. start=first page number (default 1), total=total pages (default: auto from doc)."""
    _check_rate_limit(request)
    filename = _safe_basename(file.filename)
    ext = Path(filename).suffix.lower()
    if ext not in (".pdf", ".docx", ".doc"):
        raise HTTPException(status_code=400, detail="Please upload a PDF or Word file (.pdf, .docx, .doc).")
    input_stem = Path(filename).stem or "document"

    start_num = 1
    if start:
        try:
            start_num = max(1, int(start.strip()))
        except ValueError:
            pass
    total_num = None
    if total:
        try:
            total_num = max(1, int(total.strip()))
        except ValueError:
            pass

    tmpdir = Path(tempfile.mkdtemp(prefix="convert_"))
    input_path = tmpdir / filename
    with open(input_path, "wb") as f:
        f.write(await file.read())

    try:
        total_mismatch: Optional[str] = None
        if ext == ".pdf":
            out_path, actual_pages = _add_page_numbers_pdf(input_path, start=start_num, total=total_num)
            if total_num is not None and actual_pages != total_num:
                total_mismatch = f"Document has {actual_pages} pages but you specified total={total_num}"
            media_type = "application/pdf"
            download_name = f"{input_stem}_numbered.pdf"
        else:
            if ext == ".doc":
                docx_path = tmpdir / f"{input_stem}.docx"
                _run_convert(input_path=input_path, output_path=docx_path, timeout_s=90)
                found = _find_output(tmpdir, docx_path, ".docx")
                docx_path = tmpdir / "converted.docx"
                shutil.copy2(found, docx_path)
            else:
                docx_path = input_path
            out_path = _add_page_numbers_docx(docx_path, start=start_num, total=total_num)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            download_name = f"{input_stem}_numbered.docx"
        _increment_usage(request)
        resp = _response_with_tmpdir_cleanup(out_path, media_type=media_type, download_name=download_name, tmpdir=tmpdir)
        if total_mismatch:
            resp.headers["X-Total-Mismatch"] = total_mismatch
        return resp
    except HTTPException:
        shutil.rmtree(tmpdir, ignore_errors=True)
        raise
    except Exception as e:
        shutil.rmtree(tmpdir, ignore_errors=True)
        err_msg = str(e) or "Unknown error"
        raise HTTPException(status_code=500, detail=f"Add page numbers failed: {err_msg}") from e


# ---------- REMOVE WATERMARK ----------
@app.post("/remove-watermark")
async def remove_watermark(
    request: Request,
    file: UploadFile = File(...),
):
    """Remove watermarks from PDF or Word. PDF: annotation watermarks. Word: header-based watermarks."""
    _check_rate_limit(request)
    filename = _safe_basename(file.filename)
    ext = Path(filename).suffix.lower()
    if ext not in (".pdf", ".docx", ".doc"):
        raise HTTPException(status_code=400, detail="Please upload a PDF or Word file (.pdf, .docx, .doc).")
    input_stem = Path(filename).stem or "document"

    tmpdir = Path(tempfile.mkdtemp(prefix="convert_"))
    input_path = tmpdir / filename
    with open(input_path, "wb") as f:
        f.write(await file.read())

    try:
        if ext == ".pdf":
            out_path = _remove_watermark_pdf(input_path)
            media_type = "application/pdf"
            download_name = f"{input_stem}_no_watermark.pdf"
        else:
            if ext == ".doc":
                docx_path = tmpdir / f"{input_stem}.docx"
                _run_convert(input_path=input_path, output_path=docx_path, timeout_s=90)
                docx_path = _find_output(tmpdir, docx_path, ".docx")
            else:
                docx_path = input_path
            out_path = _remove_watermark_docx(docx_path)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            download_name = f"{input_stem}_no_watermark.docx"
        _increment_usage(request)
        return _response_with_tmpdir_cleanup(out_path, media_type=media_type, download_name=download_name, tmpdir=tmpdir)
    except HTTPException:
        shutil.rmtree(tmpdir, ignore_errors=True)
        raise
    except Exception as e:
        shutil.rmtree(tmpdir, ignore_errors=True)
        raise HTTPException(status_code=500, detail=str(e)) from e


# ---------- TO HTML (for editing) ----------
@app.post("/to-html")
async def to_html(file: UploadFile = File(...)):
    """
    Convert PDF or Word to HTML for editing.
    Accepts: .pdf, .docx, .doc
    Returns: JSON with { "html": "..." }
    """
    filename = _safe_basename(file.filename)
    input_stem = Path(filename).stem or "document"
    ext = Path(filename).suffix.lower()

    if ext not in (".pdf", ".docx", ".doc"):
        raise HTTPException(
            status_code=400,
            detail="Please upload a PDF or Word document (.pdf, .docx, .doc)",
        )

    tmpdir = Path(tempfile.mkdtemp(prefix="convert_"))
    input_path = tmpdir / filename

    with open(input_path, "wb") as f:
        f.write(await file.read())

    try:
        html_path = tmpdir / f"{input_stem}.html"
        _run_convert(
            input_path=input_path,
            output_path=html_path,
            timeout_s=120,
        )
        html_content = html_path.read_text(encoding="utf-8", errors="replace")

        # Extract body content for the editor (strip full document wrapper if present)
        if "<body" in html_content.lower():
            import re
            body_match = re.search(
                r"<body[^>]*>(.*?)</body>",
                html_content,
                re.DOTALL | re.IGNORECASE,
            )
            if body_match:
                html_content = body_match.group(1).strip()

        shutil.rmtree(tmpdir, ignore_errors=True)
        return JSONResponse(content={"html": html_content})
    except Exception:
        shutil.rmtree(tmpdir, ignore_errors=True)
        raise


# ---------- HTML TO PDF ----------
@app.post("/html-to-pdf")
async def html_to_pdf(
    file: UploadFile = File(...),
    target_size_mb: Optional[str] = Form(None),
):
    """Convert HTML to PDF. Optional target_size_mb reduces file to fit (e.g. '2' for 2 MB)."""
    filename = _safe_basename(file.filename)
    if not filename.lower().endswith(".html"):
        filename = "document.html"

    tmpdir = Path(tempfile.mkdtemp(prefix="convert_"))
    input_path = tmpdir / filename

    with open(input_path, "wb") as f:
        f.write(await file.read())

    input_stem = Path(filename).stem or "document"
    pdf_path = tmpdir / f"{input_stem}.pdf"

    _run_convert(
        input_path=input_path,
        output_path=pdf_path,
        timeout_s=90,
    )
    pdf_path = _find_output(tmpdir, pdf_path, ".pdf")

    if target_size_mb:
        try:
            target_bytes = int(float(target_size_mb.strip()) * 1024 * 1024)
            if target_bytes > 0 and pdf_path.stat().st_size > target_bytes:
                pdf_path = _reduce_pdf_to_size(pdf_path, target_bytes)
        except (ValueError, TypeError):
            pass

    return _response_with_tmpdir_cleanup(
        pdf_path,
        media_type="application/pdf",
        download_name=f"{input_stem}.pdf",
        tmpdir=tmpdir,
    )


# ---------- HTML TO WORD ----------
@app.post("/html-to-word")
async def html_to_word(
    file: UploadFile = File(...),
    target_size_mb: Optional[str] = Form(None),
):
    """Convert HTML to DOCX. Optional target_size_mb reduces file to fit (e.g. '2' for 2 MB)."""
    filename = _safe_basename(file.filename)
    if not filename.lower().endswith(".html"):
        filename = "document.html"

    tmpdir = Path(tempfile.mkdtemp(prefix="convert_"))
    input_path = tmpdir / filename

    with open(input_path, "wb") as f:
        f.write(await file.read())

    input_stem = Path(filename).stem or "document"
    docx_path = tmpdir / f"{input_stem}.docx"

    # LibreOffice has no direct HTML→docx filter; use HTML→ODT→DOCX
    odt_path = tmpdir / f"{input_stem}.odt"
    _run_convert(
        input_path=input_path,
        output_path=odt_path,
        timeout_s=90,
    )
    odt_path = _find_output(tmpdir, odt_path, ".odt")
    _run_convert(
        input_path=odt_path,
        output_path=docx_path,
        timeout_s=60,
    )
    docx_path = _find_output(tmpdir, docx_path, ".docx")

    if target_size_mb:
        try:
            target_bytes = int(float(target_size_mb.strip()) * 1024 * 1024)
            if target_bytes > 0 and docx_path.stat().st_size > target_bytes:
                docx_path = _reduce_docx_to_size(docx_path, target_bytes)
        except (ValueError, TypeError):
            pass

    return _response_with_tmpdir_cleanup(
        docx_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        download_name=f"{input_stem}.docx",
        tmpdir=tmpdir,
    )


# ---------- WORD → PDF ----------
@app.post("/word-to-pdf")
async def word_to_pdf(
    request: Request,
    file: UploadFile = File(...),
    target_size_mb: Optional[str] = Form(None),
    target_size_kb: Optional[str] = Form(None),
):
    _check_rate_limit(request)
    filename = _safe_basename(file.filename)
    input_stem = Path(filename).stem or "document"

    tmpdir = Path(tempfile.mkdtemp(prefix="convert_"))
    input_path = tmpdir / filename
    pdf_path = tmpdir / f"{input_stem}.pdf"

    with open(input_path, "wb") as f:
        f.write(await file.read())

    _run_convert(
        input_path=input_path,
        output_path=pdf_path,
    )
    pdf_path = _find_output(tmpdir, pdf_path, ".pdf")

    target_bytes = None
    if target_size_kb:
        try:
            target_bytes = int(float(target_size_kb.strip()) * 1024)
        except (ValueError, TypeError):
            pass
    elif target_size_mb:
        try:
            target_bytes = int(float(target_size_mb.strip()) * 1024 * 1024)
        except (ValueError, TypeError):
            pass
    if target_bytes and target_bytes > 0 and pdf_path.stat().st_size > target_bytes:
        pdf_path = _reduce_pdf_to_size(pdf_path, target_bytes)

    _increment_usage(request)
    return _response_with_tmpdir_cleanup(
        pdf_path,
        media_type="application/pdf",
        download_name=f"{input_stem}.pdf",
        tmpdir=tmpdir,
    )


# ---------- REDUCE PDF (same format) ----------
@app.post("/reduce-pdf")
async def reduce_pdf(
    request: Request,
    file: UploadFile = File(...),
    target_size_kb: Optional[str] = Form(None),
):
    """Reduce PDF file size. Accepts PDF, returns compressed PDF."""
    _check_rate_limit(request)
    filename = _safe_basename(file.filename)
    if not filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Please upload a PDF file.")
    input_stem = Path(filename).stem or "document"

    target_bytes = None
    if target_size_kb:
        try:
            target_bytes = int(float(target_size_kb.strip()) * 1024)
        except (ValueError, TypeError):
            pass

    tmpdir = Path(tempfile.mkdtemp(prefix="convert_"))
    input_path = tmpdir / filename
    with open(input_path, "wb") as f:
        f.write(await file.read())

    out_path = input_path
    if target_bytes and target_bytes > 0 and input_path.stat().st_size > target_bytes:
        out_path = _reduce_pdf_to_size(input_path, target_bytes)

    _increment_usage(request)
    return _response_with_tmpdir_cleanup(
        out_path,
        media_type="application/pdf",
        download_name=f"{input_stem}_reduced.pdf",
        tmpdir=tmpdir,
    )


# ---------- REDUCE WORD (same format) ----------
@app.post("/reduce-word")
async def reduce_word(
    request: Request,
    file: UploadFile = File(...),
    target_size_kb: Optional[str] = Form(None),
):
    """Reduce Word file size. Accepts .doc/.docx, returns compressed DOCX."""
    _check_rate_limit(request)
    filename = _safe_basename(file.filename)
    ext = Path(filename).suffix.lower()
    if ext not in (".doc", ".docx"):
        raise HTTPException(status_code=400, detail="Please upload a Word file (.doc, .docx).")
    input_stem = Path(filename).stem or "document"

    tmpdir = Path(tempfile.mkdtemp(prefix="convert_"))
    input_path = tmpdir / filename
    html_path = tmpdir / f"{input_stem}.html"
    odt_path = tmpdir / f"{input_stem}.odt"
    docx_path = tmpdir / f"{input_stem}.docx"

    with open(input_path, "wb") as f:
        f.write(await file.read())

    # Word → HTML → ODT → DOCX
    _run_convert(input_path=input_path, output_path=html_path, timeout_s=120)
    html_path = _find_output(tmpdir, html_path, ".html")
    _run_convert(input_path=html_path, output_path=odt_path, timeout_s=90)
    odt_path = _find_output(tmpdir, odt_path, ".odt")
    _run_convert(input_path=odt_path, output_path=docx_path, timeout_s=60)
    docx_path = _find_output(tmpdir, docx_path, ".docx")

    target_bytes = None
    if target_size_kb:
        try:
            target_bytes = int(float(target_size_kb.strip()) * 1024)
        except (ValueError, TypeError):
            pass
    if target_bytes and target_bytes > 0 and docx_path.stat().st_size > target_bytes:
        docx_path = _reduce_docx_to_size(docx_path, target_bytes)

    _increment_usage(request)
    return _response_with_tmpdir_cleanup(
        docx_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        download_name=f"{input_stem}_reduced.docx",
        tmpdir=tmpdir,
    )


# ---------- MERGE DOCX ----------
@app.post("/merge-docx")
async def merge_docx(request: Request, files: List[UploadFile] = File(...)):
    """Merge multiple Word documents into one. Accepts 2+ .doc/.docx files."""
    _check_rate_limit(request)
    if len(files) < 2:
        raise HTTPException(status_code=400, detail="Please upload at least 2 Word files to merge.")

    tmpdir = Path(tempfile.mkdtemp(prefix="convert_"))
    try:
        from docx import Document

        docx_paths: List[Path] = []
        for i, uf in enumerate(files):
            fn = _safe_basename(uf.filename)
            ext = Path(fn).suffix.lower()
            if ext not in (".doc", ".docx"):
                raise HTTPException(status_code=400, detail=f"File {fn} is not a Word document (.doc, .docx).")
            path = tmpdir / f"input_{i}_{fn}"
            with open(path, "wb") as f:
                f.write(await uf.read())
            if ext == ".doc":
                docx_path = tmpdir / f"input_{i}_{Path(fn).stem}.docx"
                _run_convert(input_path=path, output_path=docx_path, timeout_s=90)
                docx_path = _find_output(tmpdir, docx_path, ".docx")
            else:
                docx_path = path
            docx_paths.append(docx_path)

        merged = Document(str(docx_paths[0]))
        for p in docx_paths[1:]:
            sub_doc = Document(str(p))
            if p != docx_paths[-1]:
                sub_doc.add_page_break()
            for element in sub_doc.element.body:
                merged.element.body.append(element)
        out_path = tmpdir / "merged.docx"
        merged.save(str(out_path))

        _increment_usage(request)
        return _response_with_tmpdir_cleanup(
            out_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            download_name="merged.docx",
            tmpdir=tmpdir,
        )
    except HTTPException:
        shutil.rmtree(tmpdir, ignore_errors=True)
        raise
    except Exception as e:
        shutil.rmtree(tmpdir, ignore_errors=True)
        raise HTTPException(status_code=500, detail=str(e)) from e


# ---------- MERGE PDF ----------
@app.post("/merge-pdf")
async def merge_pdf(request: Request, files: List[UploadFile] = File(...)):
    """Merge multiple PDFs into one. Accepts 2+ PDF files."""
    _check_rate_limit(request)
    if len(files) < 2:
        raise HTTPException(status_code=400, detail="Please upload at least 2 PDF files to merge.")

    tmpdir = Path(tempfile.mkdtemp(prefix="convert_"))
    try:
        from pypdf import PdfWriter

        merger = PdfWriter()
        for i, uf in enumerate(files):
            fn = _safe_basename(uf.filename)
            if not fn.lower().endswith(".pdf"):
                raise HTTPException(status_code=400, detail=f"File {fn} is not a PDF.")
            path = tmpdir / f"input_{i}_{fn}"
            with open(path, "wb") as f:
                f.write(await uf.read())
            merger.append(str(path))

        out_path = tmpdir / "merged.pdf"
        with open(out_path, "wb") as f:
            merger.write(f)

        _increment_usage(request)
        return _response_with_tmpdir_cleanup(
            out_path,
            media_type="application/pdf",
            download_name="merged.pdf",
            tmpdir=tmpdir,
        )
    except Exception:
        shutil.rmtree(tmpdir, ignore_errors=True)
        raise


# ---------- PDF → WORD ----------
@app.post("/pdf-to-word")
async def pdf_to_word(
    request: Request,
    file: UploadFile = File(...),
    target_size_mb: Optional[str] = Form(None),
    target_size_kb: Optional[str] = Form(None),
):
    """
    Convert PDF to DOCX. Uses PDF→HTML→ODT→DOCX pipeline because LibreOffice's
    direct PDF→DOCX can open PDFs in Draw and produce PDF instead of DOCX.
    """
    _check_rate_limit(request)
    filename = _safe_basename(file.filename)
    input_stem = Path(filename).stem or "document"

    if not filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Please upload a PDF file.")

    tmpdir = Path(tempfile.mkdtemp(prefix="convert_"))
    input_path = tmpdir / filename
    html_path = tmpdir / f"{input_stem}.html"
    odt_path = tmpdir / f"{input_stem}.odt"
    docx_path = tmpdir / f"{input_stem}.docx"

    with open(input_path, "wb") as f:
        f.write(await file.read())

    # PDF → HTML (reliable); then HTML → ODT → DOCX
    _run_convert(
        input_path=input_path,
        output_path=html_path,
        timeout_s=180,
    )
    html_path = _find_output(tmpdir, html_path, ".html")
    _run_convert(
        input_path=html_path,
        output_path=odt_path,
        timeout_s=90,
    )
    odt_path = _find_output(tmpdir, odt_path, ".odt")
    _run_convert(
        input_path=odt_path,
        output_path=docx_path,
        timeout_s=60,
    )
    docx_path = _find_output(tmpdir, docx_path, ".docx")

    target_bytes = None
    if target_size_kb:
        try:
            target_bytes = int(float(target_size_kb.strip()) * 1024)
        except (ValueError, TypeError):
            pass
    elif target_size_mb:
        try:
            target_bytes = int(float(target_size_mb.strip()) * 1024 * 1024)
        except (ValueError, TypeError):
            pass
    if target_bytes and target_bytes > 0 and docx_path.stat().st_size > target_bytes:
        docx_path = _reduce_docx_to_size(docx_path, target_bytes)

    _increment_usage(request)
    return _response_with_tmpdir_cleanup(
        docx_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        download_name=f"{input_stem}.docx",
        tmpdir=tmpdir,
    )
